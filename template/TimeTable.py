from pathlib import Path

from docx import Document


BASE_DIR = Path(__file__).resolve().parent
TITLE_SOURCE = BASE_DIR / 'Project Planning & Control with Microsoft Project.md'
OUTPUT_FILE = BASE_DIR / 'MS_Project_Training_Timetable.docx'


def load_training_title():
    for line in TITLE_SOURCE.read_text(encoding='utf-8').splitlines():
        title = line.strip()
        if title:
            return title.replace('\\&', '&')
    raise ValueError(f'No title found in {TITLE_SOURCE}')


doc = Document()

doc.add_heading(load_training_title(), 0)


def to_minutes(time_text):
    hours, minutes = map(int, time_text.split(':'))
    return (hours * 60) + minutes


def format_time(total_minutes):
    hours = total_minutes // 60
    minutes = total_minutes % 60
    return f'{hours:02d}:{minutes:02d}'


def build_row(start_minutes, duration, activity):
    end_minutes = start_minutes + duration
    return (
        f'{format_time(start_minutes)} – {format_time(end_minutes)}',
        activity,
        f'{duration} min',
    ), end_minutes


def build_day_rows(plan, end_time):
    rows = []
    current = to_minutes(plan[0][0])

    for block_start, items, expected_total in plan:
        block_start_minutes = to_minutes(block_start)
        if block_start_minutes != current:
            raise ValueError(f'Block starting at {block_start} does not align with the current timeline.')

        total = 0
        for activity, duration in items:
            row, current = build_row(current, duration, activity)
            rows.append(row)
            total += duration

        if total != expected_total:
            raise ValueError(f'Block starting at {block_start} totals {total} minutes instead of {expected_total}.')

    if current != to_minutes(end_time):
        raise ValueError(f'Day ends at {format_time(current)} instead of {end_time}.')

    return rows

def add_day(title, rows):
    doc.add_heading(title, level=1)
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Time'
    hdr_cells[1].text = 'Activity'
    hdr_cells[2].text = 'Duration'
    
    for time_slot, activity, duration in rows:
        row_cells = table.add_row().cells
        row_cells[0].text = time_slot
        row_cells[1].text = activity
        row_cells[2].text = duration


day1_plan = [
    (
        '09:00',
        [
            ('Pre-Test & Course Introduction', 10),
            ('Project Management Fundamentals\n○ PMBOK Overview\n○ Roles (Owner vs Contractor)\n○ Effort vs Duration Concept', 20),
            ('Microsoft Project Environment\n○ Interface Overview\n○ Application Configuration', 15),
            ('Lab 1: Environment Setup & Exploration', 15),
        ],
        60,
    ),
    ('10:00', [('COFFEE BREAK', 15)], 15),
    (
        '10:15',
        [
            ('Initial Project Setup\n○ Create Project File\n○ Project Summary Task', 25),
            ('Lab 2: Create Project Structure', 20),
            ('Resource Definition & Configuration\n○ Resource Types (Work, Material, Cost)\n○ Rates & Availability\n○ Base Calendar', 35),
            ('Lab 3: Resource Setup', 25),
        ],
        105,
    ),
    ('12:00', [('LUNCH', 60)], 60),
    (
        '13:00',
        [
            ('Task Creation & Scheduling Logic\n○ Task Dependencies (FS, SS, FF, SF)\n○ Lead & Lag', 40),
            ('Lab 4: Build Schedule', 25),
            ('Resource Assignment & Cost Integration\n○ Resource Assignment\n○ Cost Calculation Logic', 25),
            ('Lab 5: Assign Resources', 30),
        ],
        120,
    ),
    ('15:00', [('COFFEE BREAK', 15)], 15),
    (
        '15:15',
        [
            ('Critical Path & Optimization\n○ Identify Critical Tasks\n○ Slack Analysis\n○ Crashing & Fast Tracking', 60),
            ('Lab 6: Analyze Critical Path', 45),
        ],
        105,
    ),
]

day2_plan = [
    (
        '09:00',
        [
            ('Review & Knowledge Reinforcement', 15),
            ('Scheduling Constraints\n○ Must Start/Finish On\n○ ASAP / ALAP\n○ Deadline Management', 25),
            ('Lab 7: Apply Constraints', 20),
        ],
        60,
    ),
    ('10:00', [('COFFEE BREAK', 15)], 15),
    (
        '10:15',
        [
            ('Task Details & Visualization\n○ Notes & Indicators\n○ Timescale Customization', 20),
            ('Lab 8: Customize Views', 20),
            ('Resource Optimization\n○ Overallocation Detection\n○ Resource Leveling', 35),
            ('Lab 9: Resolve Overallocations', 30),
        ],
        105,
    ),
    ('12:00', [('LUNCH', 60)], 60),
    (
        '13:00',
        [
            ('Baseline & Project Approval', 20),
            ('Lab 10: Baseline Setup', 20),
            ('Project Execution & Tracking\n○ Task Updates\n○ Variance Analysis', 45),
            ('Lab 11: Update Progress', 35),
        ],
        120,
    ),
    ('15:00', [('COFFEE BREAK', 15)], 15),
    (
        '15:15',
        [
            ('Multi-Project Management', 20),
            ('Lab 12: Multi-Project Setup', 20),
            ('Monitoring, Analysis & Reporting', 25),
            ('Lab 13: Monitoring Dashboard', 25),
            ('Final Assessment & Closing', 15),
        ],
        105,
    ),
]

day1_rows = build_day_rows(day1_plan, '17:00')
day2_rows = build_day_rows(day2_plan, '17:00')

add_day("DAY 1 – Foundations, Configuration & Scheduling", day1_rows)
add_day("DAY 2 – Intermediate Scheduling, Execution & Monitoring", day2_rows)

doc.save(OUTPUT_FILE)

str(OUTPUT_FILE)